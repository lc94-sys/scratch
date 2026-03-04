"""
Document Processor v2 - Paragraph Chunking with Table Extraction

This module handles:
1. Extracting paragraphs from Word documents
2. Extracting and formatting tables
3. Skipping images and other non-text elements
4. Creating paragraph-based chunks with metadata

Storage Strategy:
- Small batches (<100 docs): Process in memory, directly to OpenSearch
- Large batches: Save parsed JSON to S3, then process in separate step
"""

import json
import re
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

import boto3

logger = logging.getLogger(__name__)


# =============================================================================
# DATA CLASSES
# =============================================================================

@dataclass
class ParsedParagraph:
    """Represents a parsed paragraph from the document."""
    text: str
    style: str  # e.g., "Heading 1", "Normal", "List Bullet"
    index: int  # Position in document
    element_type: str = "paragraph"


@dataclass
class ParsedTable:
    """Represents a parsed table from the document."""
    headers: List[str]
    rows: List[Dict[str, str]]
    index: int  # Position in document
    text_representation: str  # Flattened text for embedding
    element_type: str = "table"


@dataclass
class ParsedDocument:
    """Complete parsed document with all elements."""
    doc_id: str
    title: str
    summary: str
    elements: List[Any]  # List of ParsedParagraph and ParsedTable
    entitlement: List[str]
    org_id: str
    tags: List[str]
    source_file: str
    parsed_at: str
    total_paragraphs: int
    total_tables: int
    skipped_images: int


@dataclass
class DocumentChunk:
    """A chunk ready for embedding and indexing."""
    chunk_id: str
    doc_id: str
    chunk_index: int
    content: str
    content_type: str  # "paragraph" or "table"
    title: str
    summary: str
    entitlement: List[str]
    org_id: str
    tags: List[str]
    source_file: str
    # These will be added during embedding
    embedding: Optional[List[float]] = None


# =============================================================================
# DOCUMENT PARSER
# =============================================================================

class DocumentParser:
    """
    Parses Word documents (.docx) extracting paragraphs and tables.
    
    Features:
    - Paragraph extraction with style detection
    - Table extraction with header detection
    - Image skipping (logged but not processed)
    - Maintains document element order
    """
    
    def __init__(self):
        self.supported_extensions = ['.docx']
    
    def parse_document(self, file_path: str, metadata: Dict) -> ParsedDocument:
        """
        Parse a Word document and extract all text content.
        
        Args:
            file_path: Path to the .docx file
            metadata: Document metadata (title, entitlement, tags, etc.)
            
        Returns:
            ParsedDocument with all extracted elements
        """
        logger.info(f"Parsing document: {file_path}")
        
        doc = Document(file_path)
        elements = []
        element_index = 0
        skipped_images = 0
        total_paragraphs = 0
        total_tables = 0
        
        # Iterate through document body in order
        # This preserves the original document structure
        for block in self._iter_block_items(doc):
            if isinstance(block, Paragraph):
                # Check if paragraph contains images
                if self._paragraph_has_image(block):
                    skipped_images += 1
                    logger.debug(f"Skipping image in paragraph at index {element_index}")
                    continue
                
                # Skip empty paragraphs
                text = block.text.strip()
                if not text:
                    continue
                
                # Extract paragraph
                parsed_para = ParsedParagraph(
                    text=text,
                    style=block.style.name if block.style else "Normal",
                    index=element_index
                )
                elements.append(parsed_para)
                element_index += 1
                total_paragraphs += 1
                
            elif isinstance(block, Table):
                # Extract table
                parsed_table = self._parse_table(block, element_index)
                if parsed_table:
                    elements.append(parsed_table)
                    element_index += 1
                    total_tables += 1
        
        # Generate doc_id from filename
        doc_id = Path(file_path).stem.replace(' ', '_').lower()
        
        parsed_doc = ParsedDocument(
            doc_id=doc_id,
            title=metadata.get('title', doc_id),
            summary=metadata.get('summary', ''),
            elements=elements,
            entitlement=metadata.get('entitlement', ['universal']),
            org_id=metadata.get('orgId', ''),
            tags=metadata.get('metadata', {}).get('tags', []),
            source_file=Path(file_path).name,
            parsed_at=datetime.utcnow().isoformat(),
            total_paragraphs=total_paragraphs,
            total_tables=total_tables,
            skipped_images=skipped_images
        )
        
        logger.info(
            f"Parsed {file_path}: {total_paragraphs} paragraphs, "
            f"{total_tables} tables, {skipped_images} images skipped"
        )
        
        return parsed_doc
    
    def _iter_block_items(self, doc: Document):
        """
        Iterate through document blocks (paragraphs and tables) in order.
        
        This is necessary because doc.paragraphs and doc.tables are separate
        lists and don't preserve the original document order.
        """
        parent = doc.element.body
        
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)
    
    def _paragraph_has_image(self, paragraph: Paragraph) -> bool:
        """
        Check if a paragraph contains an image.
        
        Images in Word are embedded as drawings or pictures within paragraphs.
        """
        # Check for drawing elements (modern Word images)
        drawings = paragraph._element.findall('.//' + qn('w:drawing'))
        if drawings:
            return True
        
        # Check for picture elements (older format)
        pictures = paragraph._element.findall('.//' + qn('w:pict'))
        if pictures:
            return True
        
        return False
    
    def _parse_table(self, table: Table, index: int) -> Optional[ParsedTable]:
        """
        Parse a Word table into structured format.
        
        Args:
            table: python-docx Table object
            index: Position in document
            
        Returns:
            ParsedTable or None if table is empty
        """
        if len(table.rows) < 1:
            return None
        
        # First row is headers
        headers = []
        for cell in table.rows[0].cells:
            header_text = cell.text.strip()
            # Handle merged cells (may have duplicate text)
            if header_text and header_text not in headers:
                headers.append(header_text)
            elif not header_text:
                headers.append(f"Column_{len(headers) + 1}")
        
        # Remaining rows are data
        rows = []
        for row in table.rows[1:]:
            row_data = {}
            cell_texts = [cell.text.strip() for cell in row.cells]
            
            # Skip empty rows
            if not any(cell_texts):
                continue
            
            for i, cell_text in enumerate(cell_texts):
                if i < len(headers):
                    row_data[headers[i]] = cell_text
            
            if row_data:
                rows.append(row_data)
        
        if not rows:
            return None
        
        # Create text representation for embedding
        text_representation = self._table_to_text(headers, rows)
        
        return ParsedTable(
            headers=headers,
            rows=rows,
            index=index,
            text_representation=text_representation
        )
    
    def _table_to_text(self, headers: List[str], rows: List[Dict[str, str]]) -> str:
        """
        Convert table to natural language text for embedding.
        
        This creates a readable text version that captures the table's meaning.
        """
        lines = []
        
        # Add header context
        lines.append(f"Table with columns: {', '.join(headers)}")
        lines.append("")
        
        # Convert each row to a sentence
        for row in rows:
            row_parts = []
            for header, value in row.items():
                if value:
                    row_parts.append(f"{header}: {value}")
            if row_parts:
                lines.append(" | ".join(row_parts))
        
        return "\n".join(lines)


# =============================================================================
# PARAGRAPH CHUNKER
# =============================================================================

class ParagraphChunker:
    """
    Creates chunks based on paragraph boundaries.
    
    Chunking Strategy:
    - Each paragraph becomes a chunk (if within size limits)
    - Small consecutive paragraphs are merged
    - Large paragraphs are split at sentence boundaries
    - Tables become single chunks
    - Headings are prepended to following content for context
    """
    
    def __init__(
        self,
        min_chunk_size: int = 100,
        max_chunk_size: int = 1500,
        merge_small_paragraphs: bool = True,
        include_heading_context: bool = True
    ):
        """
        Initialize the chunker.
        
        Args:
            min_chunk_size: Minimum characters per chunk (merge if smaller)
            max_chunk_size: Maximum characters per chunk (split if larger)
            merge_small_paragraphs: Whether to merge consecutive small paragraphs
            include_heading_context: Whether to prepend headings to chunks
        """
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.merge_small_paragraphs = merge_small_paragraphs
        self.include_heading_context = include_heading_context
    
    def chunk_document(self, parsed_doc: ParsedDocument) -> List[DocumentChunk]:
        """
        Convert parsed document into chunks ready for embedding.
        
        Args:
            parsed_doc: Output from DocumentParser
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        chunk_index = 0
        current_heading = ""
        pending_content = []
        pending_size = 0
        
        for element in parsed_doc.elements:
            if isinstance(element, ParsedParagraph):
                # Check if this is a heading
                if self._is_heading(element.style):
                    # Flush pending content before new section
                    if pending_content:
                        chunk = self._create_chunk(
                            content="\n\n".join(pending_content),
                            content_type="paragraph",
                            chunk_index=chunk_index,
                            parsed_doc=parsed_doc,
                            heading_context=current_heading
                        )
                        chunks.append(chunk)
                        chunk_index += 1
                        pending_content = []
                        pending_size = 0
                    
                    # Update current heading for context
                    current_heading = element.text
                    
                    # Don't create a chunk for just the heading
                    # It will be prepended to the next chunk
                    continue
                
                # Regular paragraph
                para_size = len(element.text)
                
                # If paragraph is too large, split it
                if para_size > self.max_chunk_size:
                    # Flush pending first
                    if pending_content:
                        chunk = self._create_chunk(
                            content="\n\n".join(pending_content),
                            content_type="paragraph",
                            chunk_index=chunk_index,
                            parsed_doc=parsed_doc,
                            heading_context=current_heading
                        )
                        chunks.append(chunk)
                        chunk_index += 1
                        pending_content = []
                        pending_size = 0
                    
                    # Split large paragraph
                    split_chunks = self._split_large_paragraph(
                        element.text, 
                        chunk_index, 
                        parsed_doc,
                        current_heading
                    )
                    chunks.extend(split_chunks)
                    chunk_index += len(split_chunks)
                    continue
                
                # Check if we should merge with pending
                if self.merge_small_paragraphs:
                    if pending_size + para_size <= self.max_chunk_size:
                        pending_content.append(element.text)
                        pending_size += para_size
                    else:
                        # Flush pending and start new
                        if pending_content:
                            chunk = self._create_chunk(
                                content="\n\n".join(pending_content),
                                content_type="paragraph",
                                chunk_index=chunk_index,
                                parsed_doc=parsed_doc,
                                heading_context=current_heading
                            )
                            chunks.append(chunk)
                            chunk_index += 1
                        
                        pending_content = [element.text]
                        pending_size = para_size
                else:
                    # No merging - each paragraph is a chunk
                    if para_size >= self.min_chunk_size:
                        chunk = self._create_chunk(
                            content=element.text,
                            content_type="paragraph",
                            chunk_index=chunk_index,
                            parsed_doc=parsed_doc,
                            heading_context=current_heading
                        )
                        chunks.append(chunk)
                        chunk_index += 1
            
            elif isinstance(element, ParsedTable):
                # Flush pending content first
                if pending_content:
                    chunk = self._create_chunk(
                        content="\n\n".join(pending_content),
                        content_type="paragraph",
                        chunk_index=chunk_index,
                        parsed_doc=parsed_doc,
                        heading_context=current_heading
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                    pending_content = []
                    pending_size = 0
                
                # Create table chunk
                chunk = self._create_chunk(
                    content=element.text_representation,
                    content_type="table",
                    chunk_index=chunk_index,
                    parsed_doc=parsed_doc,
                    heading_context=current_heading,
                    table_data={
                        "headers": element.headers,
                        "rows": element.rows
                    }
                )
                chunks.append(chunk)
                chunk_index += 1
        
        # Don't forget remaining pending content
        if pending_content:
            chunk = self._create_chunk(
                content="\n\n".join(pending_content),
                content_type="paragraph",
                chunk_index=chunk_index,
                parsed_doc=parsed_doc,
                heading_context=current_heading
            )
            chunks.append(chunk)
        
        logger.info(f"Created {len(chunks)} chunks from document {parsed_doc.doc_id}")
        return chunks
    
    def _is_heading(self, style: str) -> bool:
        """Check if paragraph style is a heading."""
        heading_patterns = ['heading', 'title', 'toc']
        return any(pattern in style.lower() for pattern in heading_patterns)
    
    def _create_chunk(
        self,
        content: str,
        content_type: str,
        chunk_index: int,
        parsed_doc: ParsedDocument,
        heading_context: str = "",
        table_data: Optional[Dict] = None
    ) -> DocumentChunk:
        """Create a DocumentChunk with all metadata."""
        
        # Prepend heading context if available
        if self.include_heading_context and heading_context:
            content = f"{heading_context}\n\n{content}"
        
        chunk_id = f"{parsed_doc.doc_id}_chunk_{chunk_index:04d}"
        
        return DocumentChunk(
            chunk_id=chunk_id,
            doc_id=parsed_doc.doc_id,
            chunk_index=chunk_index,
            content=content,
            content_type=content_type,
            title=parsed_doc.title,
            summary=parsed_doc.summary,
            entitlement=parsed_doc.entitlement,
            org_id=parsed_doc.org_id,
            tags=parsed_doc.tags,
            source_file=parsed_doc.source_file
        )
    
    def _split_large_paragraph(
        self,
        text: str,
        start_index: int,
        parsed_doc: ParsedDocument,
        heading_context: str
    ) -> List[DocumentChunk]:
        """Split a large paragraph at sentence boundaries."""
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = start_index
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size <= self.max_chunk_size:
                current_chunk.append(sentence)
                current_size += sentence_size
            else:
                # Save current chunk
                if current_chunk:
                    chunk = self._create_chunk(
                        content=" ".join(current_chunk),
                        content_type="paragraph",
                        chunk_index=chunk_index,
                        parsed_doc=parsed_doc,
                        heading_context=heading_context if chunk_index == start_index else ""
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                current_chunk = [sentence]
                current_size = sentence_size
        
        # Don't forget the last chunk
        if current_chunk:
            chunk = self._create_chunk(
                content=" ".join(current_chunk),
                content_type="paragraph",
                chunk_index=chunk_index,
                parsed_doc=parsed_doc,
                heading_context=heading_context if chunk_index == start_index else ""
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting (handles common cases)
        sentence_endings = re.compile(r'(?<=[.!?])\s+')
        sentences = sentence_endings.split(text)
        return [s.strip() for s in sentences if s.strip()]


# =============================================================================
# STORAGE HANDLERS
# =============================================================================

class S3StorageHandler:
    """
    Handles saving/loading parsed documents to/from S3.
    
    Use this for:
    - Large batch processing
    - Production pipelines where resumability is important
    - Debugging (inspect intermediate outputs)
    """
    
    def __init__(self, bucket: str, prefix: str = "processed/"):
        self.bucket = bucket
        self.prefix = prefix
        self.s3_client = boto3.client('s3')
    
    def save_parsed_document(self, parsed_doc: ParsedDocument) -> str:
        """
        Save parsed document to S3 as JSON.
        
        Returns:
            S3 key where document was saved
        """
        # Convert to dict for JSON serialization
        doc_dict = self._parsed_doc_to_dict(parsed_doc)
        
        key = f"{self.prefix}{parsed_doc.doc_id}_parsed.json"
        
        self.s3_client.put_object(
            Bucket=self.bucket,
            Key=key,
            Body=json.dumps(doc_dict, indent=2),
            ContentType='application/json'
        )
        
        logger.info(f"Saved parsed document to s3://{self.bucket}/{key}")
        return key
    
    def load_parsed_document(self, doc_id: str) -> ParsedDocument:
        """Load parsed document from S3."""
        key = f"{self.prefix}{doc_id}_parsed.json"
        
        response = self.s3_client.get_object(Bucket=self.bucket, Key=key)
        doc_dict = json.loads(response['Body'].read())
        
        return self._dict_to_parsed_doc(doc_dict)
    
    def save_chunks(self, chunks: List[DocumentChunk], doc_id: str) -> str:
        """Save chunks to S3 as JSON."""
        chunks_dict = [self._chunk_to_dict(c) for c in chunks]
        
        key = f"{self.prefix}{doc_id}_chunks.json"
        
        self.s3_client.put_object(
            Bucket=self.bucket,
            Key=key,
            Body=json.dumps(chunks_dict, indent=2),
            ContentType='application/json'
        )
        
        logger.info(f"Saved {len(chunks)} chunks to s3://{self.bucket}/{key}")
        return key
    
    def _parsed_doc_to_dict(self, parsed_doc: ParsedDocument) -> Dict:
        """Convert ParsedDocument to JSON-serializable dict."""
        doc_dict = asdict(parsed_doc)
        
        # Convert elements to dicts
        elements = []
        for elem in parsed_doc.elements:
            if isinstance(elem, ParsedParagraph):
                elements.append({
                    "type": "paragraph",
                    "text": elem.text,
                    "style": elem.style,
                    "index": elem.index
                })
            elif isinstance(elem, ParsedTable):
                elements.append({
                    "type": "table",
                    "headers": elem.headers,
                    "rows": elem.rows,
                    "index": elem.index,
                    "text_representation": elem.text_representation
                })
        
        doc_dict['elements'] = elements
        return doc_dict
    
    def _dict_to_parsed_doc(self, doc_dict: Dict) -> ParsedDocument:
        """Convert dict back to ParsedDocument."""
        elements = []
        for elem in doc_dict['elements']:
            if elem['type'] == 'paragraph':
                elements.append(ParsedParagraph(
                    text=elem['text'],
                    style=elem['style'],
                    index=elem['index']
                ))
            elif elem['type'] == 'table':
                elements.append(ParsedTable(
                    headers=elem['headers'],
                    rows=elem['rows'],
                    index=elem['index'],
                    text_representation=elem['text_representation']
                ))
        
        doc_dict['elements'] = elements
        return ParsedDocument(**doc_dict)
    
    def _chunk_to_dict(self, chunk: DocumentChunk) -> Dict:
        """Convert DocumentChunk to dict."""
        return asdict(chunk)


class MemoryStorageHandler:
    """
    In-memory storage for small batch processing.
    
    Use this for:
    - Small batches (<100 documents)
    - Real-time processing
    - Testing/development
    """
    
    def __init__(self):
        self.parsed_documents: Dict[str, ParsedDocument] = {}
        self.chunks: Dict[str, List[DocumentChunk]] = {}
    
    def save_parsed_document(self, parsed_doc: ParsedDocument) -> str:
        """Save parsed document to memory."""
        self.parsed_documents[parsed_doc.doc_id] = parsed_doc
        return parsed_doc.doc_id
    
    def load_parsed_document(self, doc_id: str) -> ParsedDocument:
        """Load parsed document from memory."""
        return self.parsed_documents[doc_id]
    
    def save_chunks(self, chunks: List[DocumentChunk], doc_id: str) -> str:
        """Save chunks to memory."""
        self.chunks[doc_id] = chunks
        return doc_id
    
    def get_all_chunks(self) -> List[DocumentChunk]:
        """Get all chunks from all documents."""
        all_chunks = []
        for doc_chunks in self.chunks.values():
            all_chunks.extend(doc_chunks)
        return all_chunks


# =============================================================================
# MAIN PROCESSING PIPELINE
# =============================================================================

class DocumentProcessingPipeline:
    """
    Complete pipeline for processing documents from .docx to indexed chunks.
    
    Usage:
        # For small batches (in-memory)
        pipeline = DocumentProcessingPipeline(storage_type="memory")
        
        # For large batches (S3 intermediate)
        pipeline = DocumentProcessingPipeline(
            storage_type="s3",
            s3_bucket="my-bucket",
            s3_prefix="processed/"
        )
        
        # Process documents
        chunks = pipeline.process_documents(
            file_paths=["doc1.docx", "doc2.docx"],
            metadata_list=[{...}, {...}]
        )
    """
    
    def __init__(
        self,
        storage_type: str = "memory",  # "memory" or "s3"
        s3_bucket: str = None,
        s3_prefix: str = "processed/",
        min_chunk_size: int = 100,
        max_chunk_size: int = 1500
    ):
        self.parser = DocumentParser()
        self.chunker = ParagraphChunker(
            min_chunk_size=min_chunk_size,
            max_chunk_size=max_chunk_size
        )
        
        if storage_type == "s3":
            if not s3_bucket:
                raise ValueError("s3_bucket required for S3 storage")
            self.storage = S3StorageHandler(s3_bucket, s3_prefix)
        else:
            self.storage = MemoryStorageHandler()
        
        self.storage_type = storage_type
    
    def process_documents(
        self,
        file_paths: List[str],
        metadata_list: List[Dict]
    ) -> List[DocumentChunk]:
        """
        Process multiple documents through the complete pipeline.
        
        Args:
            file_paths: List of paths to .docx files
            metadata_list: List of metadata dicts (one per file)
            
        Returns:
            List of all DocumentChunks ready for embedding
        """
        all_chunks = []
        
        for file_path, metadata in zip(file_paths, metadata_list):
            try:
                # Step 1: Parse document
                parsed_doc = self.parser.parse_document(file_path, metadata)
                
                # Step 2: Save parsed document (for debugging/resumability)
                self.storage.save_parsed_document(parsed_doc)
                
                # Step 3: Create chunks
                chunks = self.chunker.chunk_document(parsed_doc)
                
                # Step 4: Save chunks
                self.storage.save_chunks(chunks, parsed_doc.doc_id)
                
                all_chunks.extend(chunks)
                
                logger.info(
                    f"Processed {file_path}: "
                    f"{parsed_doc.total_paragraphs} paragraphs, "
                    f"{parsed_doc.total_tables} tables, "
                    f"{len(chunks)} chunks created"
                )
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                raise
        
        return all_chunks
    
    def process_single_document(
        self,
        file_path: str,
        metadata: Dict
    ) -> Tuple[ParsedDocument, List[DocumentChunk]]:
        """
        Process a single document and return both parsed doc and chunks.
        
        Useful for debugging or when you need access to intermediate results.
        """
        parsed_doc = self.parser.parse_document(file_path, metadata)
        self.storage.save_parsed_document(parsed_doc)
        
        chunks = self.chunker.chunk_document(parsed_doc)
        self.storage.save_chunks(chunks, parsed_doc.doc_id)
        
        return parsed_doc, chunks


# =============================================================================
# EXAMPLE USAGE
# =============================================================================

if __name__ == "__main__":
    # Example: Process documents with in-memory storage
    pipeline = DocumentProcessingPipeline(storage_type="memory")
    
    # Example metadata
    metadata = {
        "title": "Cancellation Policy",
        "summary": "Guidelines for processing cancellation requests",
        "entitlement": ["agent_support", "agent_manager"],
        "orgId": "org_acme",
        "metadata": {
            "tags": ["cancellation", "refund", "policy"]
        }
    }
    
    # Process a single document
    # parsed_doc, chunks = pipeline.process_single_document(
    #     "cancellation_policy.docx",
    #     metadata
    # )
    
    # Print results
    # for chunk in chunks:
    #     print(f"\n--- Chunk {chunk.chunk_index} ({chunk.content_type}) ---")
    #     print(chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content)
