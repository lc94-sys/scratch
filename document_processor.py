"""
Document Processing Service

Extracts text and tables from DOCX files.
"""

import json
import os
from pathlib import Path
from typing import List, Dict
import pandas as pd
from docx import Document


class DocumentProcessor:
    """
    Processes DOCX files to extract text and tables.
    """
    
    def __init__(self, raw_docs_path: str, processed_docs_path: str):
        self.raw_docs_path = raw_docs_path
        self.processed_docs_path = processed_docs_path
        
        Path(self.processed_docs_path).mkdir(parents=True, exist_ok=True)
    
    def extract_from_docx(self, file_path: str) -> Dict:
        """Extract text and tables from DOCX"""
        doc = Document(file_path)
        extracted = {'text': '', 'tables': []}
        
        for para in doc.paragraphs:
            extracted['text'] += para.text + '\n'
        
        for table in doc.tables:
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            
            if data and len(data) > 1:
                df = pd.DataFrame(data[1:], columns=data[0])
                extracted['tables'].append(df.to_dict(orient='records'))
            
        return extracted
    
    def process_document(self, file_path: str, metadata: Dict) -> Dict:
        """Process a single DOCX document"""
        extracted = self.extract_from_docx(file_path)
        
        table_text = ""
        for i, table_dict in enumerate(extracted['tables']):
            df = pd.DataFrame(table_dict)
            table_text += f"\nTable {i+1}:\n{df.to_string(index=False)}\n"
        
        return {
            'doc_id': Path(file_path).stem,
            'content': extracted['text'] + table_text,
            'tables': extracted['tables'],
            'metadata': metadata
        }
    
    def process_all(self, documents_metadata: List[Dict]) -> List[Dict]:
        """Process all documents based on metadata list"""
        processed_documents = []
        
        for doc_meta in documents_metadata:
            file_path = os.path.join(self.raw_docs_path, doc_meta['file'])
            
            if os.path.exists(file_path):
                processed = self.process_document(
                    file_path, 
                    {
                        'entitlement': doc_meta['entitlement'],
                        'metadata': doc_meta['metadata'],
                        'orgId': doc_meta['orgId'],
                        'title': doc_meta['title'],
                        'summary': doc_meta['summary']
                    }
                )
                processed_documents.append(processed)
                
                output_file = os.path.join(
                    self.processed_docs_path, 
                    f"{processed['doc_id']}_processed.json"
                )
                with open(output_file, 'w') as f:
                    json.dump(processed, f, indent=2)
                
                print(f"✓ Processed: {doc_meta['file']}")
        
        print(f"\nTotal documents processed: {len(processed_documents)}")
        return processed_documents
    
    def load_processed_documents(self) -> List[Dict]:
        """Load all previously processed documents"""
        documents = []
        
        if not os.path.exists(self.processed_docs_path):
            return documents
        
        for filename in os.listdir(self.processed_docs_path):
            if filename.endswith('_processed.json'):
                with open(os.path.join(self.processed_docs_path, filename), 'r') as f:
                    documents.append(json.load(f))
        
        return documents
